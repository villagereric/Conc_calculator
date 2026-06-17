import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# === 計算函數 ===
def calculate_conc(c1, v2, *c2s, drug_name=''):
    data = {
        '濃度 (uM)': [],
        '需要取stock體積 (uL)': [],
        '需要加medium體積 (uL)': []
    }
    for c2 in c2s:
        v1 = (c2 * v2) / c1
        distribute = v2 - v1
        data['濃度 (uM)'].append(c2)
        data['需要取stock體積 (uL)'].append(round(v1, 2))
        data['需要加medium體積 (uL)'].append(round(distribute, 2))
    df = pd.DataFrame(data)
    return drug_name, c1, df

def calculate_combination_conc(c1_drug1, c1_drug2, v2, *c2s, drug1_name='', drug2_name=''):
    data = {
        f'{drug1_name} 濃度 (uM)': [],
        f'{drug2_name} 濃度 (uM)': [],
        f'需要取 {drug1_name} stock體積 (uL)': [],
        f'需要取 {drug2_name} stock體積 (uL)': [],
        '需要加medium體積 (uL)': []
    }
    for c2_drug1, c2_drug2 in c2s:
        v1_drug1 = (c2_drug1 * v2) / c1_drug1
        v1_drug2 = (c2_drug2 * v2) / c1_drug2
        total_drug_vol = v1_drug1 + v1_drug2
        distribute = v2 - total_drug_vol
        data[f'{drug1_name} 濃度 (uM)'].append(c2_drug1)
        data[f'{drug2_name} 濃度 (uM)'].append(c2_drug2)
        data[f'需要取 {drug1_name} stock體積 (uL)'].append(round(v1_drug1, 2))
        data[f'需要取 {drug2_name} stock體積 (uL)'].append(round(v1_drug2, 2))
        data['需要加medium體積 (uL)'].append(round(distribute, 2))
    df = pd.DataFrame(data)
    return drug1_name, c1_drug1, drug2_name, c1_drug2, df

def save_concentration_to_docx(filename, combination_result=None, single_results=None):
    doc = Document()
    if combination_result:
        drug1_name, c1_drug1, drug2_name, c1_drug2, df_combination = combination_result
        doc.add_heading(f'Combination stock_conc: {drug1_name} {c1_drug1} uM + {drug2_name} {c1_drug2} uM', level=2)
        table = doc.add_table(rows=1, cols=len(df_combination.columns))
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df_combination.columns):
            hdr_cells[i].text = col_name
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)
        for _, row in df_combination.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)
        doc.add_paragraph()

    if single_results:
        for drug_name, c1, df_single in single_results:
            doc.add_heading(f'{drug_name} stock_conc: {c1} uM', level=2)
            table = doc.add_table(rows=1, cols=len(df_single.columns))
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df_single.columns):
                hdr_cells[i].text = col_name
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)
            for _, row in df_single.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[i].paragraphs[0].runs[0].font.size = Pt(10.5)
            doc.add_paragraph()

    doc.save(filename)
    print(f"✅ Word檔案已建立完成！檔名是：{filename}")

# === GUI 執行邏輯 ===
def run_user_input():
    try:
        drug1_name = entry_d1_name.get()
        drug1_c1 = float(entry_d1_c1.get())
        drug1_c2s = list(map(float, entry_d1_c2s.get().split(',')))
        volume1 = float(entry_volume1.get())

        drug2_name = entry_d2_name.get()
        drug2_c1 = float(entry_d2_c1.get())
        drug2_c2s = list(map(float, entry_d2_c2s.get().split(',')))
        volume2 = float(entry_volume2.get())

        d1_result = calculate_conc(drug1_c1, volume1, *drug1_c2s, drug_name=drug1_name)
        d2_result = calculate_conc(drug2_c1, volume2, *drug2_c2s, drug_name=drug2_name)
        single_results = [d1_result, d2_result]

        combo_input = entry_combo.get("1.0", tk.END).strip().split('\n')
        combos = []
        for line in combo_input:
            parts = line.strip().split(',')
            if len(parts) == 2:
                combos.append((float(parts[0]), float(parts[1])))

        combination_result = None
        if combos:
            combination_result = calculate_combination_conc(
                drug1_c1, drug2_c1, volume2, *combos,
                drug1_name=drug1_name, drug2_name=drug2_name
            )

        filename = filedialog.asksaveasfilename(defaultextension=".docx",
                                                filetypes=[("Word files", "*.docx")])
        if filename:
            save_concentration_to_docx(filename,
                                       combination_result=combination_result,
                                       single_results=single_results)
            messagebox.showinfo("完成", f"Word 已儲存：{filename}")
    except Exception as e:
        messagebox.showerror("錯誤", str(e))

# === GUI 建立 ===
root = tk.Tk()
root.title("藥物濃度計算工具")
root.geometry("600x650")

# 藥物1
tk.Label(root, text="藥物1 名稱").grid(row=0, column=0, padx=5, pady=5, sticky="w")
entry_d1_name = tk.Entry(root)
entry_d1_name.insert(0, "Honokiol")
entry_d1_name.grid(row=0, column=1, padx=5, pady=5)

tk.Label(root, text="藥物1 stock濃度 (uM)").grid(row=1, column=0, padx=5, pady=5, sticky="w")
entry_d1_c1 = tk.Entry(root)
entry_d1_c1.insert(0, "50000")
entry_d1_c1.grid(row=1, column=1, padx=5, pady=5)

tk.Label(root, text="藥物1 目標濃度 (以逗號分隔)").grid(row=2, column=0, padx=5, pady=5, sticky="w")
entry_d1_c2s = tk.Entry(root, width=40)
entry_d1_c2s.insert(0, "0,5,10,15,20,25,50,75")
entry_d1_c2s.grid(row=2, column=1, padx=5, pady=5)

tk.Label(root, text="藥物1 最終體積 uL").grid(row=3, column=0, padx=5, pady=(5,15), sticky="w")
entry_volume1 = tk.Entry(root)
entry_volume1.insert(0, "2000")
entry_volume1.grid(row=3, column=1, padx=5, pady=(5,15))

# 藥物2
tk.Label(root, text="藥物2 名稱").grid(row=4, column=0, padx=5, pady=5, sticky="w")
entry_d2_name = tk.Entry(root)
entry_d2_name.insert(0, "AB")
entry_d2_name.grid(row=4, column=1, padx=5, pady=5)

tk.Label(root, text="藥物2 stock濃度 (uM)").grid(row=5, column=0, padx=5, pady=5, sticky="w")
entry_d2_c1 = tk.Entry(root)
entry_d2_c1.insert(0, "50000")
entry_d2_c1.grid(row=5, column=1, padx=5, pady=5)

tk.Label(root, text="藥物2 目標濃度 (以逗號分隔)").grid(row=6, column=0, padx=5, pady=5, sticky="w")
entry_d2_c2s = tk.Entry(root, width=40)
entry_d2_c2s.insert(0, "0,5,10,15,20,25,50,75")
entry_d2_c2s.grid(row=6, column=1, padx=5, pady=5)

tk.Label(root, text="藥物2 最終體積 uL").grid(row=7, column=0, padx=5, pady=(5,15), sticky="w")
entry_volume2 = tk.Entry(root)
entry_volume2.insert(0, "2000")
entry_volume2.grid(row=7, column=1, padx=5, pady=(5,15))

# 組合濃度
tk.Label(root, text="組合濃度 (每行一組，例如 50,100)").grid(row=8, column=0, columnspan=2, padx=5, pady=(0,5), sticky="w")
entry_combo = tk.Text(root, height=5, width=40)
entry_combo.insert(tk.END, "50,100\n50,150")
entry_combo.grid(row=9, column=0, columnspan=2, padx=5, pady=(0,15))

# 執行按鈕
btn = tk.Button(root, text="執行並儲存", command=run_user_input)
btn.grid(row=10, column=0, columnspan=2, pady=10)

root.mainloop()
